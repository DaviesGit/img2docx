from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION
from PIL import Image
import os
import sys
import shutil
import zipfile

odd_page_flag = False


def set_section_page(section, width, height):
    section.top_margin = 0
    section.bottom_margin = 0
    section.left_margin = 0
    section.right_margin = 0
    section.header_distance = 0
    section.page_width = Pt(width)
    section.page_height = Pt(height)


def set_header_img(header, file):
    header.is_linked_to_previous = False
    run = header.paragraphs[0].add_run()
    run.add_picture(file)


def change_section_img(section, file):
    width = 0
    height = 0
    with Image.open(file) as img:
        width, height = img.size
    if width <= 0 or height <= 0:
        return False
    set_section_page(section, width, height)
    set_header_img(section.header, file)
    return True


def add_img(document, file):
    global odd_page_flag
    width = 0
    height = 0
    with Image.open(file) as img:
        width, height = img.size
    if width <= 0 or height <= 0:
        return False
    new_section = document.add_section(WD_SECTION.ODD_PAGE if odd_page_flag else WD_SECTION.EVEN_PAGE)
    odd_page_flag = not odd_page_flag
    set_section_page(new_section, width, height)
    set_header_img(new_section.header, file)
    return True




def correct_img_pos(output_file):
    parent_directory = os.path.dirname(output_file)
    temp_directory_to_extract_to = (parent_directory + '/' if parent_directory else '') + os.path.splitext(
        os.path.basename(output_file))[0]
    with zipfile.ZipFile(output_file, 'r') as zip_ref:
        zip_ref.extractall(temp_directory_to_extract_to)
    index = 1
    header_file_name = temp_directory_to_extract_to + '/word/header' + str(index) + '.xml'
    while os.path.isfile(header_file_name):
        with open(header_file_name, 'r') as file:
            filedata = file.read()
        filedata = filedata.replace(
            '<wp:inline xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><wp:extent cx="21005800" cy="29705300"/><wp:docPr id="1" name="Picture 1"/><wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr><a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic>',
            '<wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251658240" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1"><wp:simplePos x="0" y="0"/><wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH><wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV><wp:extent cx="21005800" cy="29705300"/><wp:effectExtent l="0" t="0" r="6350" b="12700"/><wp:wrapNone/><wp:docPr id="1" name="Picture 1"/><wp:cNvGraphicFramePr><a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/></wp:cNvGraphicFramePr><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">')
        filedata = filedata.replace('</wp:inline></w:drawing></w:r></w:p></w:hdr>',
                                    '</wp:anchor></w:drawing></w:r></w:p></w:hdr>')
        with open(header_file_name, 'w') as file:
            file.write(filedata)

        index += 1
        header_file_name = temp_directory_to_extract_to + '/word/header' + str(index) + '.xml'

    zipf = zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED)
    for root, dirs, files in os.walk(temp_directory_to_extract_to):
        for file in files:
            zipf.write(os.path.join(root, file), os.path.join(root[len(temp_directory_to_extract_to) + 1:], file))
    zipf.close()
    shutil.rmtree(temp_directory_to_extract_to)


if len(sys.argv) < 3:
    print('The arguments number is not correct!')
    exit(1)

document = Document()
change_section_img(document.sections[0], sys.argv[1])
for i in range(2, len(sys.argv) - 1):
    add_img(document, sys.argv[i])

output_file = sys.argv[-1]
document.save(output_file)

correct_img_pos(output_file)
